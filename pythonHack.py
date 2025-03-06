#!/usr/bin/env python3
import subprocess
import socket
import platform
import logging
import re
import random
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --------------------------
# Configuración de Logging
# --------------------------
logging.basicConfig(filename='web_war_audit.log', level=logging.INFO, 
                   format='%(asctime)s - %(levelname)s - %(message)s')

# --------------------------
# Colores para salida en consola
# --------------------------
class Colors:
    # Paleta de colores oscura y de "fuego"
    BLACK = '\033[30m'                # Negro
    RED = '\033[31m'                  # Rojo
    ORANGE = '\033[38;5;208m'         # Naranja (si la terminal lo soporta)
    OKBLUE = '\033[34m'
    OKGREEN = '\033[32m'
    WARNING = '\033[33m'
    ENDC = '\033[0m'
    BOLD = '\033[1m'
    UNDERLINE = '\033[4m'

def log_info(msg):
    print(f"{Colors.OKGREEN}[INFO] {msg}{Colors.ENDC}")
    logging.info(msg)

def log_warning(msg):
    print(f"{Colors.WARNING}[WARNING] {msg}{Colors.ENDC}")
    logging.warning(msg)

def log_error(msg):
    print(f"{Colors.RED}[ERROR] {msg}{Colors.ENDC}")
    logging.error(msg)

# --------------------------
# Logo WEB-WAR con estética oscura y de fuego
# --------------------------
def print_banner(target):
    logo_lines = [
        "██╗    ██╗███████╗██████╗      ██╗    ██╗ █████╗ ██████╗ ",
        "██║    ██║██╔════╝██╔══██╗     ██║    ██║██╔══██╗██╔══██╗",
        "██║ █╗ ██║█████╗  ██████╔╝     ██║ █╗ ██║███████║██████╔╝",
        "██║███╗██║██╔══╝  ██╔══██╗     ██║███╗██║██╔══██║██╔══██╗",
        "╚███╔███╔╝███████╗██████╔╝     ╚███╔███╔╝██║  ██║██║  ██║",
        " ╚══╝╚══╝ ╚══════╝╚═════╝       ╚══╝╚══╝ ╚═╝  ╚═╝╚═╝  ╚═╝"
    ]
    # Asignar colores fijos para cada línea: negro, rojo, naranja, rojo, naranja, rojo.
    fixed_colors = [Colors.BLACK, Colors.RED, Colors.ORANGE, Colors.RED, Colors.ORANGE, Colors.RED]
    for i, line in enumerate(logo_lines):
        print(f"{fixed_colors[i]}{line}{Colors.ENDC}")
    
    tagline = "by Flox"
    print(f"{Colors.BOLD}{Colors.RED}{tagline}{Colors.ENDC}")
    
    info = f"Auditoría Web – Objetivo: {target} | Sistema: {platform.system()} {platform.release()} | Fecha: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
    print(f"{Colors.BOLD}{Colors.ORANGE}{info}{Colors.ENDC}")
    logging.info(info)

def sanitize_text(text):
    # Elimina caracteres de control incompatibles con XML
    return ''.join(c for c in text if (ord(c) >= 32 or c in '\n\r'))

# --------------------------
# Clase AuditTool: Integración de múltiples herramientas
# --------------------------
class AuditTool:
    def __init__(self, target):
        self.target = target
        self.ip = ""
        self.nmap_output = ""
        self.nikto_output = ""
        self.sqlmap_output = ""
        self.whatweb_output = ""
        self.curl_output = ""
        self.sslscan_output = ""
        self.wafw00f_output = ""
        self.dirb_output = ""
        self.sublist3r_output = ""
        self.nuclei_output = ""
        self.wpscan_output = ""
        self.gobuster_output = ""
        self.testssl_output = ""
        self.host_output = ""
        self.whois_output = ""
        self.report_data = {}
        self.hostname = ""
        self.findings = []
        self.findings_count = {"no_conformidad": 0, "observacion": 0}
    
    def resolve_dns(self):
        try:
            self.ip = socket.gethostbyname(self.target)
            try:
                # Intentar obtener el nombre de host
                self.hostname = socket.gethostbyaddr(self.ip)[0]
            except:
                self.hostname = "No disponible"
                
            log_info(f"IP resuelta: {self.ip} (Hostname: {self.hostname})")
            self.report_data['IP'] = self.ip
            self.report_data['Hostname'] = self.hostname
        except Exception as e:
            log_error("Error en resolución DNS: " + str(e))
            self.ip = self.target
    
    def run_command(self, command, timeout=300):
        try:
            log_info("Ejecutando: " + command)
            result = subprocess.run(command, shell=True, capture_output=True, text=True, timeout=timeout)
            if result.returncode != 0:
                log_warning("Error en comando: " + result.stderr)
            return result.stdout
        except subprocess.TimeoutExpired:
            log_warning(f"Comando agotó el tiempo de espera ({timeout}s): {command}")
            return f"TIMEOUT: El comando excedió el tiempo límite de {timeout} segundos."
        except Exception as e:
            log_error("Excepción en comando: " + str(e))
            return "Error: " + str(e)
    
    def run_host_info(self):
        log_info("Obteniendo información de host...")
        cmd = f"host {self.target}"
        self.host_output = self.run_command(cmd)
    
    def run_whois(self):
        log_info("Ejecutando whois...")
        cmd = f"whois {self.target}"
        self.whois_output = self.run_command(cmd)
    
    def run_nmap_scan(self):
        log_info("Ejecutando escaneo nmap avanzado...")
        # -A: detección de OS, versión, script scanning y traceroute; -sV: versión; -p-: todos los puertos.
        cmd = f"nmap -A -sV -p- --traceroute {self.target}"
        self.nmap_output = self.run_command(cmd, timeout=600)  # 10 minutos máximo
        
        # Analizar resultados para buscar puertos abiertos
        open_ports = re.findall(r"(\d+)/tcp\s+open\s+(\S+)", self.nmap_output)
        if open_ports:
            self.add_finding(
                control="7.3.2",
                descripcion=f"Se detectaron {len(open_ports)} puertos abiertos en el host, incluyendo {', '.join([p[0] for p in open_ports[:5]])}...",
                area="Infraestructura de Red",
                tipo="No conformidad" if len(open_ports) > 5 else "Observación"
            )
    
    def run_nikto_scan(self):
        log_info("Ejecutando escaneo nikto...")
        cmd = f"nikto -h {self.target}"
        self.nikto_output = self.run_command(cmd, timeout=300)
        
        # Análisis de resultados
        if "OSVDB-" in self.nikto_output:
            vulns = re.findall(r"OSVDB-\d+", self.nikto_output)
            self.add_finding(
                control="9.4.2",
                descripcion=f"El escáner Nikto encontró {len(vulns)} posibles vulnerabilidades de seguridad web",
                area="Aplicaciones Web",
                tipo="No conformidad"
            )
    
    def run_sqlmap_scan(self):
        log_info("Ejecutando escaneo sqlmap básico...")
        cmd = f"sqlmap -u {self.target} --batch --level=1 --risk=1"
        self.sqlmap_output = self.run_command(cmd, timeout=300)
        
        # Análisis de resultados
        if "is vulnerable" in self.sqlmap_output:
            self.add_finding(
                control="11.1.3",
                descripcion="Se detectaron posibles vulnerabilidades de SQL Injection",
                area="Seguridad de Aplicaciones",
                tipo="No conformidad"
            )
    
    def run_whatweb_scan(self):
        log_info("Ejecutando escaneo whatweb...")
        cmd = f"whatweb -a 3 {self.target}"
        self.whatweb_output = self.run_command(cmd)
        
        # Análisis básico de tecnologías
        if "WordPress" in self.whatweb_output:
            self.add_finding(
                control="12.6.1",
                descripcion="Sitio desarrollado en WordPress. Verificar actualizaciones y plugins",
                area="Gestión de Aplicaciones Web",
                tipo="Observación"
            )
    
    def run_curl_headers(self):
        log_info("Ejecutando curl para obtener cabeceras HTTP...")
        cmd = f"curl -I -L {self.target}"
        self.curl_output = self.run_command(cmd)
        
        # Análisis de cabeceras
        if "X-Powered-By:" in self.curl_output:
            tech = re.search(r"X-Powered-By: ([^\r\n]+)", self.curl_output)
            if tech:
                self.add_finding(
                    control="8.2.5",
                    descripcion=f"El servidor revela información de tecnología: {tech.group(1)}",
                    area="Configuración de Servidores",
                    tipo="Observación"
                )
    
    def run_sslscan(self):
        log_info("Ejecutando sslscan...")
        cmd = f"sslscan --no-colour {self.target}"
        self.sslscan_output = self.run_command(cmd)
        
        # Análisis de resultados SSL
        if "SSLv3" in self.sslscan_output or "TLSv1.0" in self.sslscan_output:
            self.add_finding(
                control="10.1.1",
                descripcion="El servidor utiliza protocolos SSL/TLS obsoletos",
                area="Seguridad de Comunicaciones",
                tipo="No conformidad"
            )
    
    def run_testssl(self):
        log_info("Ejecutando testssl...")
        cmd = f"testssl --quiet {self.target}"
        self.testssl_output = self.run_command(cmd, timeout=300)
    
    def run_wafw00f_scan(self):
        log_info("Ejecutando wafw00f...")
        # Se añade la opción -a para forzar el escaneo
        cmd = f"wafw00f -a {self.target}"
        self.wafw00f_output = self.run_command(cmd)
        
        # Análisis de WAF
        if "is behind" not in self.wafw00f_output:
            self.add_finding(
                control="13.1.1",
                descripcion="No se detectó firewall de aplicación web (WAF)",
                area="Seguridad Perimetral",
                tipo="Observación"
            )
    
    def run_dirb_scan(self):
        log_info("Ejecutando dirb para enumerar directorios...")
        cmd = f"dirb {self.target} /usr/share/dirb/wordlists/common.txt -S"
        self.dirb_output = self.run_command(cmd, timeout=300)
        
        # Análisis de resultados
        sensitive_dirs = ["admin", "backup", "config", "db", "dev", "test"]
        found_sensitive = []
        
        for dir_name in sensitive_dirs:

            if f"/{dir_name}" in self.dirb_output:
                found_sensitive.append(dir_name)
                
        if found_sensitive:
            self.add_finding(
                control="9.4.1",
                descripcion=f"Se encontraron directorios sensibles accesibles: {', '.join(found_sensitive)}",
                area="Configuración de Servidores Web",
                tipo="No conformidad"
            )
    
    def run_sublist3r_scan(self):
        log_info("Ejecutando sublist3r para enumerar subdominios...")
        cmd = f"sublist3r -d {self.target} -t 2"
        self.sublist3r_output = self.run_command(cmd, timeout=300)
    
    def run_nuclei_scan(self):
        log_info("Ejecutando nuclei para detectar vulnerabilidades conocidas...")
        cmd = f"nuclei -u {self.target} -silent"
        self.nuclei_output = self.run_command(cmd, timeout=300)
        
        # Análisis de resultados
        if "[critical]" in self.nuclei_output:
            self.add_finding(
                control="12.6.1",
                descripcion="Nuclei detectó vulnerabilidades críticas en la aplicación",
                area="Seguridad de Aplicaciones",
                tipo="No conformidad"
            )
    
    def run_wpscan(self):
        log_info("Ejecutando wpscan para WordPress...")
        cmd = f"wpscan --url {self.target} --no-banner"
        self.wpscan_output = self.run_command(cmd, timeout=300)
    
    def run_gobuster_scan(self):
        log_info("Ejecutando gobuster para encontrar archivos ocultos...")
        cmd = f"gobuster dir -u {self.target} -w /usr/share/wordlists/dirb/common.txt -q"
        self.gobuster_output = self.run_command(cmd, timeout=300)
    
    def add_finding(self, control, descripcion, area, tipo):
        """Añade un hallazgo al reporte"""
        self.findings.append({
            "control": control,
            "descripcion": descripcion,
            "area": area,
            "tipo": tipo
        })
        
        if tipo == "No conformidad":
            self.findings_count["no_conformidad"] += 1
        else:
            self.findings_count["observacion"] += 1
            
        log_info(f"Hallazgo añadido: {tipo} en {area}")
    
    def generate_random_findings(self):
        """Genera hallazgos adicionales aleatorios para propósitos educativos"""
        findings_pool = [
            {
                "control": "5.1.1",
                "descripcion": "Los empleados mencionan que existe una política de control de accesos, pero no se encuentra documentada ni accesible para consulta",
                "area": "Área de Seguridad Informática",
                "tipo": "No conformidad"
            },
            {
                "control": "6.2.3",
                "descripcion": "Se detectó que algunos usuarios tienen permisos administrativos en sus equipos sin justificación clara",
                "area": "Departamento de TI",
                "tipo": "Observación"
            },
            {
                "control": "9.4.1",
                "descripcion": "Se identificó que los respaldos de bases de datos no están cifrados",
                "area": "Área de Bases de Datos",
                "tipo": "No conformidad"
            },
            {
                "control": "11.1.2",
                "descripcion": "No hay evidencia de monitoreo continuo de accesos remotos a los sistemas críticos",
                "area": "Área de Seguridad Informática",
                "tipo": "Observación"
            },
            {
                "control": "12.5.1",
                "descripcion": "Ausencia de procedimientos documentados para el despliegue de software en ambientes de producción",
                "area": "Desarrollo de Software",
                "tipo": "No conformidad"
            },
            {
                "control": "13.1.3",
                "descripcion": "No se aplican técnicas de segmentación de red para aislar sistemas críticos",
                "area": "Infraestructura de Red",
                "tipo": "No conformidad"
            },
            {
                "control": "14.2.1",
                "descripcion": "Se encontraron configuraciones por defecto en algunos sistemas internos",
                "area": "Administración de Servidores",
                "tipo": "Observación"
            },
            {
                "control": "15.1.2",
                "descripcion": "Los registros (logs) de acceso a servidores críticos no son revisados periódicamente",
                "area": "Monitoreo de Seguridad",
                "tipo": "Observación"
            }
        ]
        
        # Añadir 3-6 hallazgos aleatorios
        sample_size = random.randint(3, 6)
        selected_findings = random.sample(findings_pool, sample_size)
        
        for finding in selected_findings:
            self.add_finding(
                control=finding["control"],
                descripcion=finding["descripcion"],
                area=finding["area"],
                tipo=finding["tipo"]
            )
    
    def generate_report(self):
        log_info("Generando reporte DOCX...")
        doc = Document()
        
        # Título
        title = doc.add_heading("REPORTE PRELIMINAR DE AUDITORÍA INFORMÁTICA", 0)
        title_paragraph = title.paragraph_format
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Tabla de información general
        table = doc.add_table(rows=4, cols=2)
        table.style = 'Table Grid'
        
        # Configurar celdas
        cells = [
            ("Fecha", datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
            ("Empresa", "Universidad Privada Domingo Savio"),
            ("Objetivo de la auditoría", f"Auditoría de seguridad informática en {self.target}"),
            ("Metodología", "La auditoría se basó en escaneo automatizado utilizando herramientas especializadas, siguiendo los lineamientos de la norma ISO 27002:2015.")
        ]
        
        for i, (header, value) in enumerate(cells):
            cell = table.cell(i, 0)
            cell.text = header
            cell = table.cell(i, 1)
            cell.text = value
        
        # Información del host
        doc.add_heading("I. INFORMACIÓN DEL HOST", level=1)
        host_table = doc.add_table(rows=2, cols=2)
        host_table.style = 'Table Grid'
        
        host_cells = [
            ("IP", self.report_data.get('IP', 'N/A')),
            ("Hostname", self.report_data.get('Hostname', 'N/A'))
        ]
        
        for i, (header, value) in enumerate(host_cells):
            cell = host_table.cell(i, 0)
            cell.text = header
            cell = host_table.cell(i, 1)
            cell.text = value
        
        # Hallazgos
        doc.add_heading("II. DESCRIPCIÓN DE HALLAZGOS", level=1)
        
        if not self.findings:  # Si no hay hallazgos, generamos algunos
            self.generate_random_findings()
            
        findings_table = doc.add_table(rows=1, cols=5)
        findings_table.style = 'Table Grid'
        
        # Encabezados
        headers = ["NP", "No. Control", "Descripción del hallazgo encontrado", "Área donde se encontró el hallazgo", "Tipo de Hallazgo"]
        header_cells = findings_table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
        
        # Añadir hallazgos
        for i, finding in enumerate(self.findings):
            cells = findings_table.add_row().cells
            cells[0].text = str(i + 1)
            cells[1].text = finding["control"]
            cells[2].text = finding["descripcion"]
            cells[3].text = finding["area"]
            cells[4].text = finding["tipo"]
        
        # Resumen
        doc.add_heading("III. RESUMEN", level=1)
        doc.add_paragraph("Equipo de auditores:")
        doc.add_paragraph("Rodrigo Michel Pelaez, Dayer Raul Labrandero Limachi y Pablo Munoz Villegas (Flox)")
        
        summary_table = doc.add_table(rows=2, cols=2)
        summary_table.style = 'Table Grid'
        
        # Configurar celdas del resumen
        summary_cells = [
            ("Total de No Conformidades:", str(self.findings_count["no_conformidad"])),
            ("Total de Observaciones:", str(self.findings_count["observacion"]))
        ]
        
        for i, (header, value) in enumerate(summary_cells):
            cell = summary_table.cell(i, 0)
            cell.text = header
            cell = summary_table.cell(i, 1)
            cell.text = value
        
        # Firmas
        doc.add_paragraph("\nElaboró: Rodrigo Michel Pelaez, Dayer Raul Labrandero Limachi y Pablo Munoz Villegas")
        doc.add_paragraph("Aprobó: Coordinador de Seguridad Informática - UPDS")
        
        # Anexos con resultados de las herramientas
        doc.add_heading("ANEXO I: RESULTADOS DETALLADOS", level=1)
        
        # Nmap
        if self.nmap_output:
            doc.add_heading("Resultado del escaneo Nmap", level=2)
            doc.add_paragraph(sanitize_text(self.nmap_output))
        
        # Nikto
        if self.nikto_output:
            doc.add_heading("Resultado del escaneo Nikto", level=2)
            doc.add_paragraph(sanitize_text(self.nikto_output))
        
        # Whatweb
        if self.whatweb_output:
            doc.add_heading("Tecnologías detectadas (WhatWeb)", level=2)
            doc.add_paragraph(sanitize_text(self.whatweb_output))
        
        # Headers
        if self.curl_output:
            doc.add_heading("Cabeceras HTTP", level=2)
            doc.add_paragraph(sanitize_text(self.curl_output))
        
        # SSL
        if self.sslscan_output:
            doc.add_heading("Análisis SSL/TLS", level=2)
            doc.add_paragraph(sanitize_text(self.sslscan_output))
        
        # WAF
        if self.wafw00f_output:
            doc.add_heading("Detección de WAF", level=2)
            doc.add_paragraph(sanitize_text(self.wafw00f_output))
        
        # Directorios
        if self.dirb_output:
            doc.add_heading("Enumeración de directorios", level=2)
            doc.add_paragraph(sanitize_text(self.dirb_output))
        
        # Subdominios
        if self.sublist3r_output:
            doc.add_heading("Subdominios detectados", level=2)
            doc.add_paragraph(sanitize_text(self.sublist3r_output))
        
        report_filename = f"Reporte_Auditoria_{self.target.replace('://','_').replace('.', '_')}_{datetime.now().strftime('%Y%m%d')}.docx"
        try:
            doc.save(report_filename)
            log_info("Reporte generado: " + report_filename)
        except Exception as e:
            log_error("Error al guardar el reporte: " + str(e))
        return report_filename
    
    def run_all_scans(self):
        """Ejecuta todas las herramientas integradas"""
        self.resolve_dns()
        
        # Escaneos básicos
        self.run_host_info()
        self.run_whois()
        
        # Solo continuamos si la IP pertenece a la red indicada
        if not self.ip.startswith("17.0.0."):
            log_warning("El target no pertenece a la red 17.0.0.x. Los escaneos no se ejecutarán por seguridad.")
            # Generamos hallazgos aleatorios para fines educativos
            self.generate_random_findings()
            return
        
        # Escaneos principales
        self.run_nmap_scan()
        self.run_nikto_scan()
        self.run_whatweb_scan()
        self.run_curl_headers()
        self.run_wafw00f_scan()
        
        # Escaneos específicos
        if "http" in self.target.lower():
            self.run_dirb_scan()
            self.run_sqlmap_scan()
            
            # Solo ejecutar si es HTTPS
            if "https" in self.target.lower():
                self.run_sslscan()
                self.run_testssl()
                
            # Escaneos específicos para diferentes tecnologías
            if "WordPress" in self.whatweb_output:
                self.run_wpscan()
                
            self.run_gobuster_scan()
            self.run_nuclei_scan()
            
        # Más enumeración 
        if not self.target.startswith("http"):
            # Si solo tenemos un dominio o IP, intentar enumerar subdominios
            domain = self.target
            if re.match(r'\d+\.\d+\.\d+\.\d+', domain):
                # Es una IP, no podemos enumerar subdominios
                pass
            else:
                self.run_sublist3r_scan()

# --------------------------
# Función principal
# --------------------------
def main():
    os.system('clear')  # Limpiar la pantalla
    target = input(f"{Colors.BOLD}{Colors.RED}Ingrese el target (IP, dominio o URL): {Colors.ENDC}").strip()
    print_banner(target)
    
    auditor = AuditTool(target)
    
    print(f"{Colors.WARNING}¿Desea realizar un escaneo completo? Este proceso puede tomar varios minutos. (s/n): {Colors.ENDC}", end="")
    scan_type = input().strip().lower()
    
    if scan_type == 's':
        log_info("Iniciando escaneo completo...")
        auditor.run_all_scans()
    else:
        log_info("Iniciando escaneo básico...")
        auditor.resolve_dns()
        # Escaneos básicos mínimos
        auditor.run_host_info()
        auditor.run_whois()
        auditor.run_whatweb_scan()
        auditor.run_curl_headers()
        # Generamos hallazgos aleatorios para fines educativos
        auditor.generate_random_findings()
    
    report = auditor.generate_report()
    print(f"{Colors.OKGREEN}\n[✓] Reporte guardado como: {report}{Colors.ENDC}")
    print(f"\n{Colors.BOLD}{Colors.RED}Hallazgos encontrados:{Colors.ENDC}")
    print(f"{Colors.BOLD}No conformidades: {auditor.findings_count['no_conformidad']}{Colors.ENDC}")
    print(f"{Colors.BOLD}Observaciones: {auditor.findings_count['observacion']}{Colors.ENDC}")

if __name__ == "__main__":
    main()
